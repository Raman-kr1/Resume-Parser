import re
import json
from typing import Dict, List, Optional, Tuple
import pdfplumber
import docx
from datetime import datetime
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self):
        # Regex patterns
        self.patterns = {
            'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            'phone': [
                r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{4,6}',
                r'[\+]?[0-9]{1,3}[-\s]?[(]?[0-9]{1,4}[)]?[-\s]?[0-9]{1,4}[-\s]?[0-9]{1,9}',
                r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
            ],
            'linkedin': r'linkedin\.com/in/[\w-]+/?',
            'github': r'github\.com/[\w-]+/?',
            'url': r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\$\$,]|(?:%[0-9a-fA-F][0-9a-fA-F]))+',
            'date': r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})',
            'year': r'\b(19|20)\d{2}\b'
        }
        
        # Skills database
        self.skills_db = self._load_skills_database()
        
        # Section headers
        self.section_headers = {
            'experience': ['experience', 'employment', 'work history', 'professional experience', 'career'],
            'education': ['education', 'academic', 'qualifications', 'academic background'],
            'skills': ['skills', 'technical skills', 'competencies', 'expertise', 'technologies'],
            'projects': ['projects', 'portfolio', 'personal projects'],
            'certifications': ['certifications', 'certificates', 'licenses'],
            'summary': ['summary', 'objective', 'profile', 'about me', 'professional summary']
        }
        
    def _load_skills_database(self) -> Dict[str, List[str]]:
        """Load comprehensive skills database"""
        return {
            'programming_languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 
                'go', 'rust', 'kotlin', 'swift', 'php', 'scala', 'r', 'matlab',
                'perl', 'shell', 'bash', 'powershell', 'sql', 'html', 'css'
            ],
            'web_frameworks': [
                'react', 'angular', 'vue.js', 'node.js', 'express.js', 'django', 
                'flask', 'spring', 'spring boot', 'asp.net', 'ruby on rails', 
                'laravel', 'symfony', 'fastapi', 'next.js', 'nuxt.js'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'oracle', 'sql server', 'sqlite',
                'redis', 'cassandra', 'dynamodb', 'elasticsearch', 'neo4j', 'couchdb'
            ],
            'cloud_platforms': [
                'aws', 'azure', 'google cloud', 'gcp', 'heroku', 'digitalocean',
                'alibaba cloud', 'ibm cloud', 'oracle cloud'
            ],
            'devops_tools': [
                'docker', 'kubernetes', 'jenkins', 'gitlab', 'github actions',
                'terraform', 'ansible', 'puppet', 'chef', 'circleci', 'travis ci'
            ],
            'data_science': [
                'machine learning', 'deep learning', 'tensorflow', 'pytorch', 
                'scikit-learn', 'pandas', 'numpy', 'keras', 'opencv', 'nlp',
                'computer vision', 'data analysis', 'statistics', 'tableau', 'power bi'
            ],
            'mobile': [
                'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic',
                'swift', 'objective-c', 'kotlin', 'android studio', 'xcode'
            ],
            'tools': [
                'git', 'svn', 'mercurial', 'jira', 'confluence', 'slack',
                'microsoft office', 'excel', 'powerpoint', 'word', 'outlook'
            ],
            'soft_skills': [
                'leadership', 'communication', 'teamwork', 'problem solving',
                'project management', 'agile', 'scrum', 'kanban', 'analytical',
                'critical thinking', 'time management', 'collaboration'
            ]
        }
    
    def read_file(self, file_path: str) -> str:
        """Read file based on extension"""
        file_ext = file_path.lower().split('.')[-1]
        
        try:
            if file_ext == 'pdf':
                return self._read_pdf(file_path)
            elif file_ext in ['docx', 'doc']:
                return self._read_docx(file_path)
            elif file_ext == 'txt':
                return self._read_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            raise
    
    def _read_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    
    def _read_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        text = ""
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
        
        return text
    
    def _read_txt(self, file_path: str) -> str:
        """Read plain text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract different sections from resume"""
        sections = {}
        lines = text.split('\n')
        
        current_section = 'header'
        section_text = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            section_found = False
            for section_name, keywords in self.section_headers.items():
                if any(keyword in line_lower for keyword in keywords):
                    # Save previous section
                    if section_text:
                        sections[current_section] = '\n'.join(section_text)
                    
                    current_section = section_name
                    section_text = []
                    section_found = True
                    break
            
            if not section_found:
                section_text.append(line)
        
        # Save last section
        if section_text:
            sections[current_section] = '\n'.join(section_text)
        
        return sections
    
    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information"""
        contact_info = {
            'name': self.extract_name(text),
            'email': self.extract_email(text),
            'phone': self.extract_phone(text),
            'linkedin': self.extract_linkedin(text),
            'github': self.extract_github(text),
            'website': self.extract_website(text)
        }
        return contact_info
    
    def extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name"""
        lines = text.split('\n')[:15]  # Name usually in first 15 lines
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines and lines with URLs/emails
            if not line or '@' in line or 'http' in line or 'www.' in line:
                continue
            
            # Skip lines with common resume words
            skip_words = ['resume', 'cv', 'curriculum', 'vitae', 'phone', 'email', 'address']
            if any(word in line.lower() for word in skip_words):
                continue
            
            # Check if line could be a name
            words = line.split()
            if 2 <= len(words) <= 4:
                # Check if words start with capital letters
                if all(word[0].isupper() for word in words if word and word[0].isalpha()):
                    # No numbers in name
                    if not any(char.isdigit() for char in line):
                        return line
        
        return None
    
    def extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        emails = re.findall(self.patterns['email'], text)
        return emails[0] if emails else None
    
    def extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number"""
        for pattern in self.patterns['phone']:
            phones = re.findall(pattern, text)
            for phone in phones:
                # Clean and validate phone number
                digits = re.sub(r'\D', '', phone)
                if 10 <= len(digits) <= 15:
                    return phone
        return None
    
    def extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn profile"""
        matches = re.findall(self.patterns['linkedin'], text, re.IGNORECASE)
        return f"https://{matches[0]}" if matches else None
    
    def extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub profile"""
        matches = re.findall(self.patterns['github'], text, re.IGNORECASE)
        return f"https://{matches[0]}" if matches else None
    
    def extract_website(self, text: str) -> Optional[str]:
        """Extract personal website"""
        urls = re.findall(self.patterns['url'], text)
        for url in urls:
            # Filter out social media and common domains
            if not any(domain in url.lower() for domain in ['linkedin', 'github', 'facebook', 'twitter']):
                return url
        return None
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume"""
        found_skills = set()
        text_lower = text.lower()
        
        # Search for skills from database
        for category, skills in self.skills_db.items():
            for skill in skills:
                # Use word boundaries for accurate matching
                pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills.add(skill.title() if len(skill) > 3 else skill.upper())
        
        return sorted(list(found_skills))
    
    def extract_education(self, text: str) -> List[Dict]:
        """Extract education information"""
        education = []
        sections = self.extract_sections(text)
        
        edu_text = sections.get('education', '')
        if not edu_text:
            edu_text = text  # Search in full text if no education section
        
        # Degree patterns
        degree_patterns = [
            r"(?:Bachelor|B\.?S\.?|B\.?A\.?|BSc|BA)(?:'s)?(?:\s+(?:of|in))?\s+(?:[\w\s]+)?",
            r"(?:Master|M\.?S\.?|M\.?A\.?|MSc|MA|MBA|M\.?Tech)(?:'s)?(?:\s+(?:of|in))?\s+(?:[\w\s]+)?",
            r"(?:Ph\.?D\.?|Doctorate)(?:\s+(?:of|in))?\s+(?:[\w\s]+)?",
            r"(?:Associate|Diploma|Certificate)(?:\s+(?:of|in))?\s+(?:[\w\s]+)?"
        ]
        
        for pattern in degree_patterns:
            matches = re.finditer(pattern, edu_text, re.IGNORECASE)
            for match in matches:
                                degree_text = match.group().strip()
                
                # Extract context around the degree
                start = max(0, match.start() - 200)
                end = min(len(edu_text), match.end() + 200)
                context = edu_text[start:end]
                
                edu_entry = {
                    'degree': degree_text,
                    'institution': None,
                    'year': None,
                    'gpa': None
                }
                
                # Extract institution
                inst_patterns = ['University', 'College', 'Institute', 'School', 'Academy']
                for inst_pattern in inst_patterns:
                    inst_match = re.search(rf'[\w\s]{{1,50}}{inst_pattern}[\w\s]{{0,20}}', context, re.IGNORECASE)
                    if inst_match:
                        edu_entry['institution'] = inst_match.group().strip()
                        break
                
                # Extract graduation year
                year_matches = re.findall(self.patterns['year'], context)
                if year_matches:
                    # Get the most recent year that makes sense
                    years = [int(y) for y in year_matches]
                    current_year = datetime.now().year
                    valid_years = [y for y in years if 1950 <= y <= current_year]
                    if valid_years:
                        edu_entry['year'] = str(max(valid_years))
                
                # Extract GPA
                gpa_match = re.search(r'(?:GPA|Grade|Score)[\s:]*(\d+\.?\d*)\s*(?:/\s*(\d+\.?\d*))?', context, re.IGNORECASE)
                if gpa_match:
                    edu_entry['gpa'] = gpa_match.group(1)
                
                education.append(edu_entry)
        
        # Remove duplicates
        seen = set()
        unique_education = []
        for edu in education:
            key = (edu['degree'], edu['institution'])
            if key not in seen:
                seen.add(key)
                unique_education.append(edu)
        
        return unique_education
    
    def extract_experience(self, text: str) -> List[Dict]:
        """Extract work experience"""
        experience = []
        sections = self.extract_sections(text)
        
        exp_text = sections.get('experience', '')
        if not exp_text:
            # Try to find experience section in full text
            exp_match = re.search(
                r'(?:experience|employment|work history)[\s:]*\n([\s\S]+?)(?=\n[A-Z]{2,}|education|skills|$)',
                text,
                re.IGNORECASE
            )
            if exp_match:
                exp_text = exp_match.group(1)
        
        if not exp_text:
            return []
        
        # Split into individual job entries
        # Look for patterns that indicate new job entry
        job_splits = re.split(r'\n(?=\w+.*?(?:Engineer|Developer|Manager|Analyst|Consultant|Designer|Specialist|Director|Lead|Senior|Junior))', exp_text)
        
        for job_text in job_splits:
            if not job_text.strip():
                continue
            
            lines = job_text.strip().split('\n')
            if not lines:
                continue
            
            exp_entry = {
                'position': None,
                'company': None,
                'duration': None,
                'location': None,
                'description': []
            }
            
            # First line often contains position and company
            first_line = lines[0]
            
            # Extract position
            position_keywords = ['Engineer', 'Developer', 'Manager', 'Analyst', 'Consultant', 
                               'Designer', 'Specialist', 'Director', 'Lead', 'Architect',
                               'Administrator', 'Coordinator', 'Executive', 'Officer',
                               'Scientist', 'Researcher', 'Intern', 'Associate']
            
            for keyword in position_keywords:
                if keyword in first_line:
                    # Extract position part
                    pos_match = re.search(rf'[\w\s]*{keyword}[\w\s]*', first_line, re.IGNORECASE)
                    if pos_match:
                        exp_entry['position'] = pos_match.group().strip()
                        break
            
            # Extract company (often after 'at', '@', '|', '-')
            company_match = re.search(r'(?:at|@|\|)\s*([^,\n|]+)', first_line, re.IGNORECASE)
            if company_match:
                exp_entry['company'] = company_match.group(1).strip()
            
            # Extract dates
            date_pattern = r'(?:' + self.patterns['date'] + r')\s*[-–—]\s*(?:' + self.patterns['date'] + r'|Present|Current)'
            date_match = re.search(date_pattern, job_text, re.IGNORECASE)
            if date_match:
                exp_entry['duration'] = date_match.group()
            
            # Extract location
            location_pattern = r'(?:,\s*|in\s+)([\w\s]+,\s*\w{2,})'
            location_match = re.search(location_pattern, first_line)
            if location_match:
                exp_entry['location'] = location_match.group(1).strip()
            
            # Extract bullet points (job descriptions)
            for line in lines[1:]:
                line = line.strip()
                if line and (line.startswith(('•', '-', '▪', '→')) or re.match(r'^\d+\.', line)):
                    # Clean bullet point
                    desc = re.sub(r'^[•\-▪→\d.)\s]+', '', line).strip()
                    if desc:
                        exp_entry['description'].append(desc)
                elif line and len(exp_entry['description']) < 5:  # Limit descriptions
                    exp_entry['description'].append(line)
            
            if exp_entry['position'] or exp_entry['company']:
                experience.append(exp_entry)
        
        return experience
    
    def extract_projects(self, text: str) -> List[Dict]:
        """Extract projects information"""
        projects = []
        sections = self.extract_sections(text)
        
        proj_text = sections.get('projects', '')
        if not proj_text:
            return []
        
        # Split into individual projects
        project_entries = re.split(r'\n(?=\w+.*?(?:Project|App|System|Platform|Tool|Website))', proj_text)
        
        for proj in project_entries:
            if not proj.strip():
                continue
            
            lines = proj.strip().split('\n')
            
            project = {
                'name': lines[0].strip() if lines else None,
                'description': [],
                'technologies': []
            }
            
            # Extract descriptions and technologies
            for line in lines[1:]:
                line = line.strip()
                if not line:
                    continue
                
                # Check if line mentions technologies
                if any(tech in line.lower() for tech in ['technologies:', 'tech stack:', 'built with:', 'using:']):
                    # Extract technologies
                    tech_text = line.lower()
                    for category, skills in self.skills_db.items():
                        for skill in skills:
                            if skill.lower() in tech_text:
                                project['technologies'].append(skill)
                else:
                    project['description'].append(line)
            
            if project['name']:
                projects.append(project)
        
        return projects
    
    def extract_certifications(self, text: str) -> List[Dict]:
        """Extract certifications"""
        certifications = []
        sections = self.extract_sections(text)
        
        cert_text = sections.get('certifications', '')
        if not cert_text:
            # Search in full text
            cert_match = re.search(r'(?:certifications?|certificates?)[\s:]*\n([\s\S]+?)(?=\n[A-Z]{2,}|$)', text, re.IGNORECASE)
            if cert_match:
                cert_text = cert_match.group(1)
        
        if cert_text:
            lines = cert_text.strip().split('\n')
            
            for line in lines:
                line = line.strip()
                if not line or len(line) < 5:
                    continue
                
                cert = {
                    'name': line,
                    'year': None,
                    'issuer': None
                }
                
                # Extract year
                year_match = re.search(self.patterns['year'], line)
                if year_match:
                    cert['year'] = year_match.group()
                    cert['name'] = line.replace(year_match.group(), '').strip()
                
                # Common certification issuers
                issuers = ['Microsoft', 'Google', 'Amazon', 'AWS', 'Oracle', 'Cisco', 'CompTIA', 'PMI', 'Scrum']
                for issuer in issuers:
                    if issuer.lower() in line.lower():
                        cert['issuer'] = issuer
                        break
                
                certifications.append(cert)
        
        return certifications
    
    def calculate_score(self, parsed_data: Dict) -> Dict[str, float]:
        """Calculate resume completeness score"""
        scores = {
            'contact_info': 0.0,
            'experience': 0.0,
            'education': 0.0,
            'skills': 0.0,
            'overall': 0.0
        }
        
        # Contact info score
        contact_fields = ['name', 'email', 'phone']
        contact_count = sum(1 for field in contact_fields if parsed_data.get('contact', {}).get(field))
        scores['contact_info'] = (contact_count / len(contact_fields)) * 100
        
        # Experience score
        exp_count = len(parsed_data.get('experience', []))
        scores['experience'] = min(100, exp_count * 25)  # 25 points per experience, max 100
        
        # Education score
        edu_count = len(parsed_data.get('education', []))
        scores['education'] = min(100, edu_count * 50)  # 50 points per degree, max 100
        
        # Skills score
        skills_count = len(parsed_data.get('skills', []))
        scores['skills'] = min(100, skills_count * 5)  # 5 points per skill, max 100
        
        # Overall score
        scores['overall'] = sum([
            scores['contact_info'] * 0.2,
            scores['experience'] * 0.3,
            scores['education'] * 0.2,
            scores['skills'] * 0.3
        ])
        
        return scores
    
    def parse(self, file_path: str) -> Dict:
        """Main parsing method"""
        try:
            logger.info(f"Parsing resume: {file_path}")
            
            # Read file
            text = self.read_file(file_path)
            if not text:
                return {'error': 'No text could be extracted from the file'}
            
            # Extract sections
            sections = self.extract_sections(text)
            
            # Parse resume
            parsed_data = {
                'file': file_path,
                'contact': self.extract_contact_info(text),
                'skills': self.extract_skills(text),
                'education': self.extract_education(text),
                'experience': self.extract_experience(text),
                'projects': self.extract_projects(text),
                'certifications': self.extract_certifications(text),
                'sections': list(sections.keys()),
                'text_length': len(text),
                'parsed_date': datetime.now().isoformat()
            }
            
            # Add completeness score
            parsed_data['scores'] = self.calculate_score(parsed_data)
            
            logger.info(f"Successfully parsed resume: {file_path}")
            return parsed_data
            
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {e}")
            return {'error': str(e), 'file': file_path}